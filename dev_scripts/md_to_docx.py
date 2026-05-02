from docx import Document
import re

def create_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        l = line.strip()
        if not l:
            doc.add_paragraph("")
            continue
            
        if l.startswith("### "):
            doc.add_heading(l[4:], level=3)
        elif l.startswith("## "):
            doc.add_heading(l[3:], level=2)
        elif l.startswith("# "):
            doc.add_heading(l[2:], level=1)
        elif l.startswith("* ") or l.startswith("- "):
            doc.add_paragraph(l[2:], style='List Bullet')
        else:
            # Handle basic bolding translation
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', l)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    # Italics
                    subparts = re.split(r'(\*.*?\*)', part)
                    for sub in subparts:
                        if sub.startswith('*') and sub.endswith('*') and len(sub)>2:
                            p.add_run(sub[1:-1]).italic = True
                        else:
                            p.add_run(sub)

    doc.save(docx_path)
    print("Done generating docx")

if __name__ == "__main__":
    create_docx("DriftNet_Presentation_Guide.md", "DriftNet_Presentation_Guide.docx")
